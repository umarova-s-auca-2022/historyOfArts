"""
Art History Study Guide Builder
================================
Put this file in the SAME folder as your e-course index.html
(next to all the File_xxx/ subfolders), then run:

    python make_study_guide.py

It will create:  Art_History_Study_Guide.docx

Requirements (install once):
    pip install python-docx Pillow
"""

import os, glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ──────────────────────────────────────────────────────────────────

def find_image(folder):
    """Return the first image file found inside a subfolder, or None."""
    if not os.path.isdir(folder):
        return None
    for ext in ('*.jpg','*.jpeg','*.png','*.gif','*.bmp','*.webp'):
        hits = glob.glob(os.path.join(folder, ext), recursive=False)
        if hits:
            return hits[0]
    return None

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x2e, 0x4a)
    return p

def add_artwork_title(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{num})  ")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0x1a, 0x2e, 0x4a)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x1a, 0x2e, 0x4a)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'B8860B')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_image(doc, path):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=Inches(4.5))
    except Exception as e:
        doc.add_paragraph(f"[Image could not be loaded: {e}]")

def add_qa(doc, q, a):
    # Question
    pq = doc.add_paragraph()
    pq.paragraph_format.left_indent  = Pt(18)
    pq.paragraph_format.space_before = Pt(5)
    pq.paragraph_format.space_after  = Pt(2)
    rq = pq.add_run(q)
    rq.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    # Answer
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent  = Pt(36)
    pa.paragraph_format.space_before = Pt(2)
    pa.paragraph_format.space_after  = Pt(6)
    ra = pa.add_run("Answer:  ")
    ra.bold = True
    ra.font.size = Pt(11)
    ra.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    rb = pa.add_run(a)
    rb.font.size = Pt(11)

def page_break(doc):
    doc.add_page_break()

# ── artwork data ──────────────────────────────────────────────────────────────
# Each entry: (number, title, subfolder_name, [(question, answer), ...])
# subfolder_name is the EXACT folder name inside the e-course directory.

ARTWORKS = [

# ── POTTERY ──────────────────────────────────────────────────────────────────
("ANCIENT GREEK ART — POTTERY & VASE PAINTING", None, None, None),

(1, "Funerary Crater, Dipylon, Athens, 8th c. BCE",
 "File_Funerary_crater_Dipylon_..._.179253", [
  ("Q1. Who made this and what is it called?", "An anonymous work known as the Dipylon Funerary Crater, made in Athens around the 8th century BCE. Named after the Dipylon Gate cemetery where similar vessels were found."),
  ("Q2. What period and style?", "The Geometric period (c. 900–700 BCE), specifically the Late Geometric phase."),
  ("Q3. What does it represent?", "A grave marker. Its painted decoration includes a prothesis (lying in state of the deceased) surrounded by mourning figures, and processions of chariots and warriors arranged in horizontal registers around the vessel."),
  ("Q4. What stylistic characteristics are evident?", "Geometric style: all figures are reduced to schematic silhouettes made of geometric shapes (triangles for torsos, circles for heads). Composition is organized in horizontal registers. No perspective or depth — all figures are flat and symbolic. The Meander (Greek key) pattern is a signature element."),
]),

(2, "Polyphemus Painter, Amphora, c. 650 BCE",
 "File_Polyphemus_painter_Ampho..._.179254", [
  ("Q1. Who made this and what is it?", "Attributed to the Polyphemus Painter (named after this work), c. 650 BCE. Orientalizing period."),
  ("Q2. What period and style?", "Orientalizing period (c. 700–600 BCE) — the transitional phase between Geometric and Archaic, when Greek art absorbed Near Eastern and Egyptian influences."),
  ("Q3. What does it represent?", "One side depicts Odysseus and his men blinding the Cyclops Polyphemus (from Homer's Odyssey) — one of the earliest surviving narrative mythological scenes in Greek vase painting."),
  ("Q4. What stylistic characteristics are evident?", "Orientalizing style: animal friezes, floral and vegetal motifs (rosettes, lotus flowers) borrowed from Near Eastern art. Figures are still silhouetted but more rounded than Geometric ones. Narrative mythological scenes begin to appear."),
]),

(3, "Polyphemus / Eleusis Amphora, c. 650 BCE",
 "File_Popyphemus_Eleusis_650_BCE_.179255", [
  ("Q1. What is this work?", "The Eleusis Amphora, found at Eleusis near Athens, c. 650 BCE. Attributed to the Polyphemus Painter or his circle. Proto-Attic style."),
  ("Q2. What does it represent?", "The neck shows Odysseus blinding Polyphemus; the body shows Perseus fleeing the Gorgons after beheading Medusa. Among the earliest large-scale mythological narrative compositions in Greek art."),
  ("Q3. What stylistic characteristics are evident?", "Uses outline drawing and incised details plus added white and red paint — more flexible than pure black silhouette. Figures are larger and more expressive than Geometric. Narrative approach reflects growing interest in Homeric epics."),
]),

(4, "Exekias, Achilles and Penthesilea, c. 525 BCE",
 "File_Exekias_Achiles_and_Pent..._.297783", [
  ("Q1. Who is the artist and what is the title?", "Exekias, one of the greatest Athenian black-figure painters, created this kylix (drinking cup) c. 525 BCE."),
  ("Q2. What period and technique?", "Late Archaic period, black-figure technique, c. 525 BCE."),
  ("Q3. What does it represent?", "Achilles kills Penthesilea, Queen of the Amazons. At the moment of her death their eyes met and Achilles fell in love with her — a moment of tragic beauty."),
  ("Q4. What stylistic characteristics are evident?", "Black-figure technique: figures painted in black slip on red clay; details incised (scratched) through the black. Exekias is celebrated for psychological depth — the locked gaze of the two figures conveys emotional intensity unusual for this period."),
]),

(5, "Exekias, Dionysos at Sea, c. 530 BCE",
 "File_Exekias_Dionysos_at_sea_..._.179256", [
  ("Q1. Who is the artist and what is the title?", "Exekias, kylix interior, c. 530 BCE. The Dionysus Cup. State Collections of Antiquities, Munich."),
  ("Q2. What does it represent?", "Dionysus reclines in a boat surrounded by a sprouting grapevine and leaping dolphins. Myth: pirates tried to abduct him; he turned them into dolphins."),
  ("Q3. What stylistic characteristics are evident?", "Masterpiece of black-figure painting. Exekias uses the circular tondo format brilliantly — the boat's hull echoes the cup's shape. The grapevine fills the space with graceful organic forms. Despite the flatness of black-figure, he creates a vivid sense of sea, wind, and divine magic."),
]),

(6, "Berlin Painter, Bell Crater: Abduction of Europa, 490 BCE",
 "File_Berlin_painter_bell_crat..._.179257", [
  ("Q1. Who is the artist?", "Attributed to the Berlin Painter (named after a work in Berlin), c. 490 BCE. Red-figure bell crater."),
  ("Q2. What period and technique?", "Early Classical (Late Archaic) period, red-figure technique, c. 490 BCE."),
  ("Q3. What does it represent?", "Europa, a Phoenician princess, is abducted by Zeus who has transformed into a white bull, riding across the sea toward Crete."),
  ("Q4. What stylistic characteristics are evident?", "Red-figure technique: figures left in red clay while the background is painted black; internal details drawn with a fine brush — allowing greater naturalism than black-figure. The Berlin Painter is known for placing isolated figures against plain black backgrounds without ground lines, creating a monumental quality."),
]),

(7, "Penthesilea Painter, Achilles and Penthesilea, 455 BCE",
 "File_Penthesilea_painter_Achi..._.179258", [
  ("Q1. Who is the artist?", "The Penthesilea Painter (named after this kylix), c. 455 BCE. Red-figure technique."),
  ("Q2. What does it represent?", "Same myth as Exekias: Achilles kills Penthesilea; their eyes meet at the moment of her death."),
  ("Q3. How does it compare to the Exekias version?", "Red-figure allows for detailed rendering of muscles, drapery, and facial expressions. Compared to Exekias's black-figure version, this shows greater anatomical naturalism and more dynamic movement. The Early Classical style moves away from Archaic rigidity toward more organic, expressive forms."),
]),

(8, "Sosias Painter, Achilles Tending Patroclus, 500 BCE",
 "File_Achilles_tending_Patrocl..._.179259", [
  ("Q1. Who is the artist?", "Attributed to the Sosias Painter, c. 500 BCE. Red-figure kylix interior. Antikensammlungen, Berlin."),
  ("Q2. What does it represent?", "From the Iliad: Achilles carefully bandages the wounded arm of Patroclus, who looks away in pain. A rare depiction of tenderness between two warriors."),
  ("Q3. What stylistic characteristics are evident?", "Celebrated for psychological intimacy. Advanced foreshortening of the bandaged arm. The focused gaze of Achilles and the turned-away face of Patroclus express emotion through body language — marking the transition from Archaic formalism to Classical humanism."),
]),

(9, "Niobid Painter, Kalyx Crater, c. 460 BCE",
 "File_Niobid_painter_kalyx_cra..._.179264", [
  ("Q1. Who is the artist?", "The Niobid Painter, c. 460 BCE. Red-figure calyx-krater. Louvre, Paris."),
  ("Q2. What does it represent?", "Apollo and Artemis shoot down the children of Niobe (the Niobids), who boasted she was superior to goddess Leto. Figures are scattered at different levels suggesting a landscape."),
  ("Q3. What is revolutionary about this work?", "Figures are placed at different heights on the surface to suggest depth and a landscape setting — unlike earlier vase painting where all figures stood on a single ground line. This innovation reflects the influence of large-scale mural painting. Figures in complex poses — fallen, running, dying — with considerable anatomical sophistication."),
]),

(10, "Reed Painter, Warrior by a Grave, c. 410 BCE (White-Ground)",
 "File_Reed_painter_Warrior_by_..._.179265", [
  ("Q1. Who is the artist and what technique?", "The Reed Painter, c. 410 BCE. White-ground lekythos (funerary oil flask). Late Classical period."),
  ("Q2. What does it represent?", "A warrior stands by a grave stele (grave marker). White-ground lekythoi depicted mourning scenes and were used as tomb offerings."),
  ("Q3. What stylistic characteristics are evident?", "White-ground technique: white slip surface, figures drawn in outline with added colours (red, yellow, blue). Allows a wider palette than red-figure but more fragile — used mainly for funerary items. The style is looser and more painterly, approaching actual panel painting. The quiet, melancholy mood reflects the contemplative nature of funerary art."),
]),

# ── SCULPTURE ────────────────────────────────────────────────────────────────
("ANCIENT GREEK ART — SCULPTURE", None, None, None),

(11, "New York Kouros from Attica, c. 600 BCE",
 "File_New_York_couros_from_Att..._.179266", [
  ("Q1. Who made this and what is it called?", "'Kouros' means 'young man' in Greek. The New York Kouros is anonymous, from Attica, c. 600 BCE. Metropolitan Museum of Art, New York."),
  ("Q2. What period?", "Archaic period, c. 600 BCE."),
  ("Q3. What does it represent?", "A standing nude male youth — either a grave marker or a votive offering to the gods. It may represent a god, athlete, or idealized young man."),
  ("Q4. What Archaic stylistic characteristics are evident?", "Strict frontality — the figure faces directly forward. Both arms hang rigidly at the sides with clenched fists. One foot slightly advanced (Egyptian influence), but weight is distributed equally — no real movement. The anatomy is schematic and idealized. The 'Archaic smile' (slight upward curve of the lips). Hair depicted as decorative patterned beads."),
  ("Q5. Why is it important?", "The kouros type spans two centuries; its development shows Greek sculpture's progression from rigid Archaic formalism to Classical naturalism."),
]),

(12, "Kritios Boy, c. 480 BCE",
 "File_Kritos_boy_480_BCE_.179268", [
  ("Q1. What is this work?", "Attributed to the sculptor Kritios, c. 480 BCE. Found on the Acropolis in Athens. Acropolis Museum, Athens."),
  ("Q2. What period?", "Early Classical / Severe Style, c. 480 BCE. Marks the transition from Archaic to Classical."),
  ("Q3. What makes the Kritios Boy different from the kouros?", "The Kritios Boy is the first known Greek sculpture to use contrapposto: the weight shifts to one leg, the hip drops on the free leg side, and the torso responds with a subtle counter-movement — creating a sense of natural, relaxed standing. The Archaic smile is gone — the face is serious and naturalistic. This work marks the birth of the Classical style."),
  ("Q4. How does it compare to the New York Kouros?", "New York Kouros (Archaic): rigid, frontal, schematic anatomy, Archaic smile, both legs equal. Kritios Boy (Early Classical): contrapposto, naturalistic anatomy, no smile, psychologically present. The contrast perfectly illustrates the Archaic-to-Classical transition."),
]),

(13, "Peplos Kore, c. 530 BCE",
 "File_Peplos_Kore_530_BCE_.179267", [
  ("Q1. What is this work?", "The Peplos Kore is an anonymous Archaic marble sculpture, c. 530 BCE. 'Kore' means 'maiden.' Acropolis Museum, Athens."),
  ("Q2. What does it represent?", "A standing clothed female figure, likely a votive offering to Athena on the Acropolis. She wears a peplos (heavy woolen garment)."),
  ("Q3. What Archaic characteristics are evident?", "Like the kouros: rigid frontality, weight evenly distributed, Archaic smile. The garment is stylized with decorative patterned folds rather than naturalistic cloth. Traces of paint show she was originally brightly coloured. Unlike the nude kouros, women were always depicted clothed in Greek art of this period."),
]),

(14, "Charioteer of Delphi, c. 475 BCE (bronze)",
 "File_Charioteer_.179260", [
  ("Q1. What is this work?", "The Charioteer of Delphi, rare surviving Greek bronze, c. 475 BCE. Delphi Archaeological Museum, Greece."),
  ("Q2. What period?", "Early Classical / Severe Style, c. 475 BCE."),
  ("Q3. What Severe Style characteristics are evident?", "A departure from Archaic rigidity — more naturalistic head, inlaid glass eyes giving an uncanny lifelike quality, copper lips. The pose is still relatively frontal and composed. The vertical folds of the long robe create column-like stability. A masterpiece of early bronze casting."),
]),

(15, "Warriors from Riace, c. 450 BCE (bronze)",
 "File_Warrior_from_Riace_450_.179269", [
  ("Q1. What are these works?", "Two Greek bronze warriors found off the coast of Riace, Italy, c. 450 BCE. National Museum of Magna Graecia, Reggio Calabria."),
  ("Q2. What Classical characteristics are evident?", "Perfect musculature, balanced proportions, naturalistic anatomy. Contrapposto creates natural, relaxed movement. Inlaid eyes, copper lips and nipples, silver teeth give extraordinary lifelike presence — rare survivals of original Greek bronzes."),
]),

(16, "Zeus or Poseidon, Artemision Bronze, c. 450 BCE",
 "File_Poseidon_or_Zeus_bronze_..._.179270", [
  ("Q1. What is this work?", "The Artemision Bronze (Zeus or Poseidon), c. 450 BCE, found off Cape Artemision. National Archaeological Museum, Athens."),
  ("Q2. What does it represent?", "A god — either Zeus hurling a thunderbolt or Poseidon throwing a trident (the weapon is missing). The powerful stance captures the moment of action."),
  ("Q3. What Classical characteristics are evident?", "Both arms extended horizontally, creating a wide dynamic stance visible from multiple angles. Idealized but deeply naturalistic anatomy. Balanced yet energetic — the Classical achievement of capturing movement while maintaining harmony and control."),
]),

(17, "Polykleitos, Doryphoros (Spear Bearer), c. 440 BCE",
 "File_Polykleitos_Doryphoros_4..._.179271", [
  ("Q1. Who is the artist?", "Polykleitos of Argos, c. 440 BCE. Original bronze lost; known through Roman marble copies. National Archaeological Museum, Naples."),
  ("Q2. What is its significance?", "Polykleitos wrote the 'Canon' — a theoretical treatise defining the perfect mathematical proportions of the human body (head = 1/7 of total height). The Doryphoros was a three-dimensional demonstration of this Canon."),
  ("Q3. What stylistic characteristics are evident?", "The figure embodies contrapposto perfectly: the weight-bearing right leg is tense, the free left leg is relaxed; the right arm hangs while the left holds the spear — a rhythmic opposing diagonal movement (chiasmus). This balance of opposites creates both movement and stability. The defining embodiment of Classical ideals: harmony, proportion, and beauty of the idealized male body."),
  ("Q4. How does it compare to the Kritios Boy and the Boxer at Rest?", "Kritios Boy (Early Classical): first contrapposto, naturalistic but small. Doryphoros (High Classical): fully developed contrapposto with mathematical precision, the definitive Canon. Boxer (Hellenistic): abandons mathematical perfection for psychological realism, emotional expression, and the marks of a lived life."),
]),

(18, "Myron, Discobolus (Discus Thrower), c. 450 BCE",
 "File_Myron_Discobolus_450BC_.179272", [
  ("Q1. Who is the artist?", "Myron, c. 450 BCE. Original bronze lost; known through Roman marble copies. Palazzo Massimo alle Terme, Rome."),
  ("Q2. What does it represent?", "An athlete at the moment of maximum coil before releasing the discus — the split second before the throw."),
  ("Q3. What stylistic characteristics are evident?", "Myron's genius was to capture a frozen moment of maximum dynamic potential. The composition is designed as a flat relief — most effective from one angle. Despite the dynamic pose, the athlete's face remains calm and expressionless — the Classical ideal that emotion should not disturb the beauty of the body."),
]),

(19, "Wounded Amazon, c. 430 BCE (marble copy of bronze original)",
 "File_Wounded_Amazon_430_BCE_m..._.179273", [
  ("Q1. Who made this?", "Attributed to either Polykleitos or Phidias, c. 430 BCE. Known through Roman marble copies."),
  ("Q2. What does it represent?", "An Amazon (warrior woman) who has been wounded, shown resting with one arm raised against a pillar."),
  ("Q3. What stylistic characteristics are evident?", "Significant as an early depiction of a female figure in a near-nude, athletic context. Shows Classical balance — pain is suggested by the raised arm and weight shift, but the face remains ideally composed. Drapery falls naturally from the waist. Represents an expansion of Classical sculptural subjects from purely athletic males to female heroic figures."),
]),

(20, "Praxiteles, Aphrodite of Knidos, c. 350 BCE",
 "File_Praxiteles_Aphrodite_of_..._.228295", [
  ("Q1. Who is the artist?", "Praxiteles of Athens, c. 350 BCE. Original lost; known through Roman marble copies. Vatican Museums."),
  ("Q2. What is its significance?", "The first monumental sculpture of a fully nude female figure in Greek art — a revolutionary work. Previously, female figures were always clothed."),
  ("Q3. What stylistic characteristics are evident?", "Praxiteles developed the S-curve (sinuous body curve) that creates languid grace and sensuality. The figure's pose — turning slightly, covering her body — creates movement and psychological presence. Ancient travelers came specifically to Knidos to see it. It defined the female nude in Western art for centuries."),
]),

(21, "Aphrodite of Milos (Venus de Milo), c. 100–130 BCE",
 "File_Aphrodite_of_Milos_100-1..._.179289", [
  ("Q1. What is this work?", "The Aphrodite of Milos (Venus de Milo) by Alexandros of Antioch, c. 100–130 BCE. Louvre, Paris."),
  ("Q2. What period?", "Hellenistic period, c. 100–130 BCE, combining Classical idealism with Hellenistic complexity."),
  ("Q3. What stylistic characteristics are evident?", "The sinuous S-curve and idealized facial features reflect Praxitelean influence. The dramatic twist of the torso is a Hellenistic development. The drapery loosely fallen around the hips creates a sharp contrast between the nude upper body and the clothed lower body. One of the most famous sculptures in the world."),
]),

(22, "Lysippos, Apoxyomenos (Scraper), c. 320 BCE",
 "File_Lysippos_Apoxyomenos_Rom..._.179290", [
  ("Q1. Who is the artist?", "Lysippos of Sikyon, c. 320 BCE. Known through Roman marble copies. Vatican Museums."),
  ("Q2. What does it represent?", "An athlete cleaning himself after exercise by scraping oil and sand from his skin with a strigil (metal scraper). Arms extended forward, breaking into the viewer's space."),
  ("Q3. How does Lysippos differ from Polykleitos?", "Lysippos rejected Polykleitos's Canon, using a new system: smaller head (1/8 body height instead of 1/7), longer legs, slimmer body. The outstretched arms project into the viewer's space, making the figure fully three-dimensional — it must be viewed from multiple angles, a major Late Classical innovation. He was the court sculptor of Alexander the Great."),
]),

(23, "Lysippos (after), Farnese Hercules, c. 320 BCE",
 "File_Herakles_Farmese_Glykon_..._.179292", [
  ("Q1. What is this work?", "Roman marble copy by Glykon of Athens of an original bronze by Lysippos, c. 320 BCE. National Archaeological Museum, Naples."),
  ("Q2. What does it represent?", "Hercules rests after completing his Twelve Labors — leaning heavily on his club, massive body exhausted. Behind his back he holds (hidden from the viewer) the golden apples of the Hesperides."),
  ("Q3. What is significant about this work?", "Marks the Late Classical move toward psychological and emotional expression. Unlike the confident Doryphoros, Hercules conveys weariness and vulnerability despite superhuman physical power. The contrast between his enormous body and defeated, resting posture is emotionally powerful — pointing toward Hellenistic expressiveness."),
]),

# ── HELLENISTIC ──────────────────────────────────────────────────────────────
("HELLENISTIC PERIOD SCULPTURE", None, None, None),

(24, "Head of Alexander the Great, marble, c. 330–300 BCE",
 "File_Head_of_Alexandre_the_gr..._.179293", [
  ("Q1. What does it represent?", "Alexander the Great, king of Macedonia. Typically shown with an upward-tilting gaze (the 'anastole'), idealized features, and swept-back lion's mane of hair."),
  ("Q2. What stylistic characteristics are evident?", "Alexander's portraits established a new ideal: the divine ruler-hero. The 'pathetic' style — slightly parted lips, intense upward glance toward the heavens — was a new development toward psychological expressiveness. His image blended Classical ideal beauty with individualized, emotionally expressive traits. Lysippos was his official portraitist."),
]),

(25, "Boxer at Rest (Terme Boxer), c. 100–50 BCE (bronze)",
 "File_Boxer_at_rest_bronze_330..._.179294", [
  ("Q1. Who made this and what is its title?", "The Boxer at Rest (also: Terme Boxer, Boxer of the Quirinal), an original Greek bronze, 2nd–1st century BCE. National Museum of Rome, Palazzo Massimo alle Terme."),
  ("Q2. What period?", "Hellenistic period, c. 100–50 BCE."),
  ("Q3. What does it represent?", "An elderly, heavily built boxer seated on a rock after a fight. He turns his head as if responding to someone. He wears leather boxing thongs (himantes). His face and body show a career's worth of damage — broken nose, cauliflower ears, cuts — rendered in inlaid copper."),
  ("Q4. What Hellenistic characteristics are evident?", "Hellenistic realism and psychological depth. Unlike the Classical ideal of the perfect athletic youth, this boxer is old, damaged, and vulnerable — a real, suffering individual. Inlaid copper simulates fresh blood and bruises — graphic naturalism. Hellenistic sculpture expanded subjects to include the ugly, aged, suffering, and emotionally complex."),
  ("Q5. How does it compare to the Doryphoros?", "Doryphoros (Classical): perfect youth, mathematical proportions, idealized body, calm expression — the ideal of what a human should be. Boxer (Hellenistic): old, battered, real — depicting what a human actually becomes. Classical art shows the ideal; Hellenistic art shows the truth."),
]),

(26, "Dying Gaul, c. 230–220 BCE (Roman marble copy)",
 "File_Dying_gaul_.179298", [
  ("Q1. What is this work?", "Roman marble copy of a bronze commissioned by Attalos I of Pergamon to commemorate his victory over the Gauls, c. 230–220 BCE. Capitoline Museums, Rome."),
  ("Q2. What does it represent?", "A dying Gaul (Celtic warrior) sinking to the ground after battle. His torque (Celtic neck ring) and wound identify him as a barbarian enemy. Despite being a defeated enemy, he is portrayed with dignity and sympathy."),
  ("Q3. What Hellenistic characteristics are evident?", "Interest in the suffering and dignity of non-Greeks. Rather than triumphalism, the focus is on the noble, stoic death of the foe — creating empathy. Highly naturalistic anatomy conveying physical collapse convincingly. This expansion of subject matter (non-Greek peoples, dying figures, emotional extremity) defines Hellenistic style."),
]),

(27, "Winged Nike of Samothrace, c. 190 BCE",
 "File_Winged_Nike_of_Samothrace_.179299", [
  ("Q1. What is this work?", "The Winged Nike (Victory) of Samothrace, Greek marble, c. 190 BCE. Louvre, Paris."),
  ("Q2. What does it represent?", "The goddess Nike alighting on the prow of a ship, wings spread, garments blown back by the wind. Created to commemorate a naval victory."),
  ("Q3. What Hellenistic characteristics are evident?", "The drapery shows extraordinary virtuosity — wind-blown fabric clings to and reveals the body beneath ('wet drapery') while billowing dramatically. The figure conveys tremendous energy and forward momentum. Unlike calm Classical figures, this work is all movement, action, and drama. The multi-directional pose requires multiple viewing angles."),
]),

(28, "Laocoön and His Sons, 1st c. BCE–1st c. CE",
 "File_Laocoon_and_his_two_sons..._.467286", [
  ("Q1. Who made this?", "Attributed to three Rhodian sculptors: Hagesandros, Polydoros, and Athanadoros. Vatican Museums, Rome."),
  ("Q2. What does it represent?", "The Trojan priest Laocoön and his two sons attacked and killed by sea serpents sent by the gods for warning the Trojans against the wooden horse."),
  ("Q3. What Hellenistic characteristics are evident?", "The defining example of Hellenistic Baroque style: extreme expression of suffering, movement, and emotional anguish. Figures writhe in agony, twisted in multiple planes, faces contorted with pain. Diagonal compositions, complex intertwining poses, and expression of extreme terror are key features. Enormously influential on Renaissance and Baroque artists."),
]),

# ── ARCHITECTURE ─────────────────────────────────────────────────────────────
("GREEK ARCHITECTURE & TEMPLE SCULPTURE", None, None, None),

(29, "Temple of Zeus at Olympia — Metope: Heracles and Atlas",
 "File__Heracles_and_Atlas_meto..._.183984", [
  ("Q1. What is this work?", "Marble metope relief from the Temple of Zeus at Olympia, c. 460 BCE. Early Classical (Severe Style). Archaeological Museum of Olympia."),
  ("Q2. What does it represent?", "Heracles holds up the sky (temporarily replacing Atlas) while Athena assists. Atlas returns with the golden apples of the Hesperides — one of Heracles's twelve labors."),
  ("Q3. What stylistic characteristics are evident?", "Severe Style: simple, powerful forms with minimal ornamentation. Figures are broad, heavy, and solid. Emotions are underplayed — contained power rather than dramatic expression. Athena's calm assistance contrasts with the exertion of Heracles."),
]),

(30, "Temple of Zeus at Olympia — West Pediment: Apollo, Lapith and Centaur",
 "File__Apollo_with_Lapith_and_..._.183989", [
  ("Q1. What is this work?", "West pediment sculpture of the Temple of Zeus at Olympia, c. 460 BCE. Early Classical (Severe Style)."),
  ("Q2. What does it represent?", "The Centauromachy — battle between the Lapiths and Centaurs. Apollo stands at center, arm outstretched, commanding order over the chaos of the battle."),
  ("Q3. What is the symbolic significance?", "Apollo's central, calm presence embodies Classical order, reason, and divine authority. The composition expresses the Greek cultural conflict: civilization vs. barbarism, reason vs. chaos — themes relevant to the Greek victory over the Persians."),
]),

(31, "Phidias, Statue of Zeus in the Temple of Zeus at Olympia (reconstruction)",
 "File_statue_of_Zeus_in_the_te..._.183997", [
  ("Q1. What is this work?", "The cult statue of Zeus by Phidias, c. 435 BCE, Temple of Zeus at Olympia. One of the Seven Wonders of the Ancient World. The original is lost — known only through ancient descriptions."),
  ("Q2. What did it represent?", "Zeus seated on a magnificent throne, chryselephantine (gold and ivory over a wooden core), approximately 12–13 metres tall."),
  ("Q3. What is its significance?", "Considered by the ancients one of the greatest works ever made. The scale and material luxury (gold and ivory) made it a visible expression of divine and civic power. Phidias was also responsible for the Athena Parthenos and the sculptural program of the Parthenon."),
]),

(32, "The Parthenon, Athens, 447–432 BCE",
 "File_The_Parthenon_.183986", [
  ("Q1. What is this work?", "The Parthenon, Doric temple on the Acropolis of Athens, built 447–432 BCE by architects Iktinos and Kallikrates under the supervision of Phidias."),
  ("Q2. What does it represent?", "The primary temple of Athena Parthenos (Athena the Virgin), patron goddess of Athens. A monument to Athenian power following the Persian Wars."),
  ("Q3. What are its architectural characteristics?", "The supreme achievement of Greek Doric temple architecture. Incorporates optical refinements: columns are slightly convex (entasis) to appear straight; the floor curves upward slightly; columns lean inward slightly — all counteracting optical illusions. The sculptural program included metope reliefs (Centauromachy, Amazonomachy), a continuous Ionic frieze (the Panathenaic procession), and pedimental sculptures."),
]),

(33, "Lapith and Centaur Metope, The Parthenon, c. 447–432 BCE",
 "File_Lapith_and_Centaur_metop..._.183990", [
  ("Q1. What is this work?", "High-relief metope from the south side of the Parthenon, c. 447–432 BCE. British Museum, London."),
  ("Q2. What does it represent?", "A Centaur attacks or wrestles a Lapith warrior. The Centauromachy symbolized the triumph of Greek civilization over barbarism — an allegory of Athens's victory over the Persians."),
  ("Q3. What Classical characteristics are evident?", "Ideally proportioned, anatomically superb, dynamically composed. Violent struggle conveyed through diagonal poses and intertwining figures. Despite the violence, a Classical sense of harmony and control — every form beautifully sculpted, composition balanced within the square metope field."),
]),

(34, "Phidias, Athena Parthenos (reconstruction)",
 "File_Athena_in_Pathenon_.183996", [
  ("Q1. What is this work?", "Phidias, cult statue of Athena Parthenos for the Parthenon, completed c. 438 BCE. The original is lost; known through small-scale copies and ancient descriptions."),
  ("Q2. What did it represent?", "Athena, goddess of wisdom and warfare, approximately 12 metres tall in gold and ivory (chryselephantine). She wore a full helmet and aegis, held a shield, and in her outstretched hand stood a small figure of Nike (Victory)."),
  ("Q3. What is its significance?", "Considered one of the greatest works of the Classical world. Gold and ivory materials expressed Athenian wealth and power. Phidias oversaw the entire sculptural program of the Parthenon, making him the most important artist of the Classical period."),
]),

(35, "Nike Adjusting Her Sandal, Temple of Athena Nike, c. 410 BCE",
 "File_Nike_adjusting_the_sanda..._.183992", [
  ("Q1. What is this work?", "Relief panel from the parapet of the Temple of Athena Nike on the Acropolis, c. 410 BCE. Acropolis Museum, Athens."),
  ("Q2. What does it represent?", "Nike bends over to adjust her sandal — a casual, intimate moment for a divine figure. A transient everyday gesture rather than a grand mythological act."),
  ("Q3. What Late Classical stylistic characteristics are evident?", "'Rich style': the drapery is complex, thin, and fluid, clinging to and revealing the body beneath in the 'wet drapery' effect — the garment appears almost transparent. An intimate, casual pose is a departure from the grand frontal poses of the High Classical period."),
]),

# ── EARLY RENAISSANCE ─────────────────────────────────────────────────────────
("EARLY RENAISSANCE (c. 1400–1490)", None, None, None),

(36, "Masaccio, The Expulsion from Eden, c. 1427",
 "File__Masaccio_The_Expulsion_..._.189624", [
  ("Q1. Who is the artist?", "Masaccio (Tommaso di Ser Giovanni di Simone), Brancacci Chapel frescoes, Florence, c. 1427."),
  ("Q2. What does it represent?", "Adam and Eve expelled from Paradise by an angel after eating the forbidden fruit. Eve cries out in anguish; Adam covers his face in shame."),
  ("Q3. What Early Renaissance characteristics are evident?", "Masaccio revolutionized painting: (1) Naturalistic anatomy — real 3D bodies with weight and gravity; (2) Emotional expression — Eve's open-mouthed cry of grief is unprecedented in its rawness; (3) Chiaroscuro to model form; (4) Psychological depth. A radical break from the flat, decorative figures of medieval and International Gothic style."),
]),

(37, "Masaccio, The Holy Trinity, c. 1427",
 "File_Masaccio_The_Holy_Trinity_.189626", [
  ("Q1. Who is the artist?", "Masaccio, Santa Maria Novella, Florence, c. 1427."),
  ("Q2. What does it represent?", "The Trinity: God the Father behind the Cross supports the crucified Christ; the Holy Spirit is between them. The Virgin Mary and St. John stand at the foot. Below: two kneeling donors. At the very bottom: a painted skeleton inscribed 'I was once what you are, and what I am you will be.'"),
  ("Q3. Why is this work historically significant?", "The first known use of strict mathematical linear perspective in a painted work. Masaccio used Brunelleschi's perspective system to create a trompe-l'oeil barrel vault that appeared to open the wall — contemporary viewers thought the wall had been opened."),
]),

(38, "Donatello, David, c. 1440s–1460s (bronze)",
 "File_Donatello_David_.189627", [
  ("Q1. Who is the artist?", "Donatello, bronze David, Florence, c. 1440s–1460s, for the Medici family. Bargello Museum, Florence."),
  ("Q2. What does it represent?", "David, the young biblical hero who killed Goliath, stands triumphant, his foot on Goliath's severed head. He wears a shepherd's hat and boots — otherwise nude."),
  ("Q3. What is the significance?", "The first large-scale freestanding nude sculpture since antiquity — a revolutionary act reviving Classical Greek nudity and contrapposto. Created for private Medici courtyard display — a sign of Renaissance humanism."),
  ("Q4. How does it compare to Michelangelo's David and Bernini's David?", "Donatello (Early Renaissance): calm, post-victory, androgynous, contemplative. Michelangelo (High Renaissance): tense, pre-battle, heroic and idealized. Bernini (Baroque): mid-action, dynamic twisting, grimacing with effort. Each David reflects the values and style of its period."),
]),

(39, "Donatello, Mary Magdalen, c. 1453–1455 (wood)",
 "File__Donatello_Mary_Magdalen_.189629", [
  ("Q1. Who is the artist?", "Donatello, polychrome wood, c. 1453–1455. Museo dell'Opera del Duomo, Florence."),
  ("Q2. What does it represent?", "Mary Magdalen as a penitent saint who withdrew to the desert as an ascetic. She is shown gaunt and aged, her body wasted, her hair and garments in wild matted tangles, hands clasped in prayer."),
  ("Q3. What makes this work significant?", "Unlike the idealized figures of Classical or High Renaissance art, this Magdalen is deliberately aged, emaciated, and ravaged — her asceticism physically visible. This expressive realism shows Donatello's range: he could create idealized beauty (the David) but also penetrating psychological and spiritual truth."),
]),

(40, "Botticelli, The Birth of Venus, c. 1484–1486",
 "File__Botticelli_The_birth_of..._.189631", [
  ("Q1. Who is the artist?", "Sandro Botticelli, c. 1484–1486. Uffizi Gallery, Florence."),
  ("Q2. What does it represent?", "Venus (goddess of love, born from sea foam) arrives on the shore of Cyprus on a giant scallop shell. Zephyr (wind god) blows her toward shore. The Hora (goddess of spring) rushes to clothe her."),
  ("Q3. What Early Renaissance characteristics are evident?", "Connected to the Neoplatonic circle around Lorenzo de' Medici: Venus represents divine beauty (the Neoplatonic concept of Celestial Venus). Sinuous flowing lines rather than Masaccio's weighty realism. The mythological subject (not religious) reflects humanist interest in classical antiquity. The nude female echoes Greek sculpture (Venus Pudica pose)."),
  ("Q4. What connection to ancient Greek culture does this work demonstrate?", "Direct: Venus is a goddess from the classical pantheon (Greek: Aphrodite). In terms of art: Botticelli revives Classical Greek ideals of beauty, harmony, and proportion — particularly the Venus Pudica type — and combines them with Renaissance humanism, celebrating the human body as noble and divine rather than sinful."),
]),

(41, "Botticelli, Primavera (Spring), c. 1478",
 "File_Botticelli_Primavera_.189632", [
  ("Q1. Who is the artist?", "Sandro Botticelli, c. 1478. Uffizi Gallery, Florence."),
  ("Q2. What does it represent?", "A complex allegory set in an orange grove: from right, Zephyr pursues Chloris who transforms into Flora (spring), scattering flowers; centre: Venus with Cupid above; left: three dancing Graces and Mercury dispersing clouds."),
  ("Q3. What characteristics are evident?", "A humanist allegory reflecting Neoplatonic philosophy. Botticelli favors linear grace and flowing movement over weight and volume. Figures have a dreamlike, tapestry-like quality. The complex Neoplatonic symbolism reflects the intellectual environment of Medici Florence."),
]),

(42, "Botticelli, Mars and Venus, c. 1483",
 "File__Botticelli_Mars_and_Venus_.189633", [
  ("Q1. Who is the artist?", "Sandro Botticelli, c. 1483. National Gallery, London."),
  ("Q2. What does it represent?", "Venus (goddess of love) gazes alertly at Mars (god of war) who sleeps in a deep slumber. Mischievous satyrs play with his armour. The Neoplatonic idea: Love (Venus) conquers War (Mars)."),
  ("Q3. What characteristics are evident?", "The Neoplatonic allegory enacted through the contrast between the alert, clothed Venus and the deeply sleeping nude Mars. Botticelli's characteristic linear grace is evident throughout."),
]),

(43, "Mantegna, Dead Christ (Lamentation), c. 1480",
 "File_Mantegna_Dead_Christ_.189635", [
  ("Q1. Who is the artist?", "Andrea Mantegna, c. 1480. Pinacoteca di Brera, Milan."),
  ("Q2. What does it represent?", "The dead body of Christ after the crucifixion, seen from the feet, dramatically foreshortened. Grieving figures (the Virgin, St. John) weep at the upper left."),
  ("Q3. What makes it significant?", "A tour-de-force of foreshortening — representing the body in extreme perspective as it recedes toward the viewer. Mantegna was deeply influenced by classical sculpture; Christ's body has the quality of sculpted marble. Combines humanist realism with theological meditation."),
]),

(44, "Mantegna, Ceiling of Camera degli Sposi, c. 1473",
 "File__Mantegna_The_ceiling_of..._.189636", [
  ("Q1. Who is the artist?", "Andrea Mantegna, ceiling of the Camera degli Sposi, Gonzaga Palace, Mantua, c. 1473."),
  ("Q2. What is its significance?", "The first illusionistic ceiling painting (di sotto in sù — 'from below looking up'): a painted circular oculus appears to open to the sky, with figures and putti looking down from a balustrade. Unprecedented use of extreme foreshortening (figures seen from below). Enormously influential — anticipates Michelangelo's Sistine ceiling and all Baroque ceiling painting."),
]),

(45, "Mantegna, Parnassus, 1497",
 "File_Mantegnia_Andrea_Parnassus_.232977", [
  ("Q1. Who is the artist?", "Andrea Mantegna, 1497, for Isabella d'Este's studiolo. Louvre, Paris."),
  ("Q2. What does it represent?", "On Mount Parnassus, Mars and Venus stand on a rocky arch; the nine Muses dance below; Apollo plays music; Mercury and Pegasus observe from the left. Vulcan (Venus's cuckolded husband) gestures angrily from his cave."),
  ("Q3. What characteristics are evident?", "Painted for a humanist patron as an allegory of the arts. The precise, crystalline quality of the landscape is typical of North Italian Renaissance art. The scene celebrates music, poetry, and the arts — appropriate for a private studiolo celebrating learning and culture."),
]),

# ── HIGH RENAISSANCE ─────────────────────────────────────────────────────────
("HIGH RENAISSANCE (c. 1490–1527)", None, None, None),

(46, "Leonardo da Vinci, The Last Supper, 1495–1498",
 "File_Leonardo_The_Last_Supper_.190398", [
  ("Q1. Who is the artist?", "Leonardo da Vinci, 1495–1498. Refectory of Santa Maria delle Grazie, Milan."),
  ("Q2. What does it represent?", "The final meal of Jesus with his twelve apostles (John 13:21). The moment: Jesus announces 'One of you will betray me.' The apostles react with shock, denial, and agitation. Judas (fourth from left) clutches his money bag and recoils."),
  ("Q3. What High Renaissance characteristics are evident?", "Leonardo used strict one-point perspective with the vanishing point behind Jesus's head — making Christ the mathematical and visual centre. The apostles are arranged in four groups of three, each expressing a different emotional reaction. Embodies High Renaissance ideals: perfect harmony, rational order, psychological depth, and ideal beauty."),
  ("Q4. How does it compare to Tintoretto's Last Supper (Baroque)?", "Leonardo (High Renaissance, pre-Reformation): calm, symmetrical, horizontal, all figures clearly visible, rational clarity — humanist art confident in reason. Tintoretto (Baroque, post-Counter-Reformation): diagonal table rushing into deep space, chaotic atmosphere, ordinary servants mixed with disciples, supernatural angels materialising in smoke — emotional overwhelm rather than rational clarity."),
]),

(47, "Leonardo da Vinci, Mona Lisa, c. 1503–1519",
 "File_Leonardo_Mona_Lisa_.190399", [
  ("Q1. Who is the artist?", "Leonardo da Vinci, c. 1503–1519. Louvre, Paris."),
  ("Q2. What techniques does Leonardo use? (Masters of Illusion)", "(1) Sfumato: subtle blurring of edges and tonal transitions, creating the soft atmospheric quality and the mysterious smile that appears to change depending on where you look; (2) Atmospheric perspective: the landscape becomes progressively bluer, hazier, and less detailed in the distance, creating convincing depth. The three-quarter pose (vs. profile) was a Renaissance innovation."),
]),

(48, "Leonardo da Vinci, Madonna of the Rocks, c. 1483–1486",
 "File_Leonardo_Madonna_of_the_..._.469348", [
  ("Q1. Who is the artist?", "Leonardo da Vinci, c. 1483–1486. Louvre version (Paris); a later version is in the National Gallery, London."),
  ("Q2. What does it represent?", "The Virgin Mary shelters the infant Jesus and infant John the Baptist in a mysterious rocky grotto. An angel points toward John."),
  ("Q3. What Renaissance techniques are evident?", "The pyramid composition of the four figures (Mary at apex) shows Leonardo's compositional mastery. Sfumato throughout — figures emerge from the shadowy grotto with soft edges. Meticulous scientific observation of geology in the rock formations."),
]),

(49, "Leonardo da Vinci, Madonna and Child with Saint Anne, c. 1503–1519",
 "File__Leonardo_Madonna_and_Ch..._.190400", [
  ("Q1. Who is the artist?", "Leonardo da Vinci, c. 1503–1519. Louvre, Paris."),
  ("Q2. What does it represent?", "Saint Anne sits with the Virgin Mary on her lap; Mary leans forward to restrain the infant Christ who reaches toward a lamb (symbol of his future sacrifice). Three generations in an intertwined pyramid composition."),
  ("Q3. What High Renaissance characteristics are evident?", "The pyramidal grouping is the defining compositional formula of the High Renaissance — creating stability, harmony, and natural togetherness. Leonardo uses sfumato extensively. The tender psychological interaction between the three figures reflects Leonardo's deep interest in human emotion and relationship."),
]),

(50, "Michelangelo, David, 1501–1504 (marble)",
 "File_Michelangelo_David_.190402", [
  ("Q1. Who is the artist?", "Michelangelo Buonarroti, 1501–1504. Galleria dell'Accademia, Florence."),
  ("Q2. What does it represent?", "David before his battle with Goliath — the moment before the fight. David is alert, tense, assessing his enemy, his sling over his shoulder. Unlike Donatello's post-battle David, this is the pre-battle moment of decision."),
  ("Q3. What High Renaissance characteristics are evident?", "Michelangelo revives and surpasses Classical ideals: idealized yet intensely naturalistic anatomy, unprecedented mastery of musculature. Contrapposto recalls the Doryphoros but the psychological tension (taut focused gaze, swollen neck vein, clenched fist) is entirely Renaissance. Over 5 metres tall — a monument of civic and humanist pride."),
]),

(51, "Michelangelo, Pietà, 1498–1499 (marble)",
 "File__Michelangelo_Pieta_.190403", [
  ("Q1. Who is the artist?", "Michelangelo Buonarroti, 1498–1499. St. Peter's Basilica, Vatican City. He was 23–24 when he completed it. The only work he ever signed."),
  ("Q2. What does it represent?", "The Virgin Mary holds the dead body of Christ after the crucifixion. 'Pietà' (Italian: 'pity')."),
  ("Q3. What characteristics are evident?", "Michelangelo solved the compositional problem of a seated woman holding an adult man's limp body by spreading Mary's garments into a broad pyramidal base. Mary appears remarkably young. Technical virtuosity: drapery carved with hairlike delicacy; Christ's body idealized and serene. Classical idealization combined with Christian devotion."),
]),

(52, "Michelangelo, The Creation of Adam (Sistine Chapel), 1508–1512",
 "File__Michelangelo_The_creati..._.190404", [
  ("Q1. Who is the artist?", "Michelangelo, Sistine Chapel ceiling, Vatican, 1508–1512."),
  ("Q2. What does it represent?", "God gives life to Adam: God, surrounded by angels, reaches toward the reclining Adam. Their fingers almost touch — the moment before the infusion of the divine spark."),
  ("Q3. What characteristics are evident?", "God depicted as a powerful, dynamic elderly patriarch hurtling through space; Adam as a perfect Classical male nude, languidly reclining. The almost-touching fingertips is the most famous image in Western art. Michelangelo humanizes the divine and divinizes the human — a quintessentially Renaissance (Neoplatonic) idea."),
]),

(53, "Michelangelo, The Last Judgment, 1536–1541 (Sistine Chapel)",
 "File__Michelangelo_The_Last_J..._.190405", [
  ("Q1. Who is the artist?", "Michelangelo, altar wall of the Sistine Chapel, Vatican, 1536–1541."),
  ("Q2. What does it represent?", "The Second Coming of Christ and final judgment: Christ raises his arm in condemnation. To his left, the damned fall into Hell; to his right, the saved rise to Heaven. St. Bartholomew holds his own flayed skin — traditionally seen as Michelangelo's self-portrait."),
  ("Q3. How does it differ from the Sistine ceiling?", "Painted decades later, this work reflects a shift toward Mannerism and Counter-Reformation spirituality. Figures are larger, more muscular, and more agitated; the composition is turbulent and centrifugal rather than rationally ordered. Christ is muscular and terrifying rather than gentle."),
]),

(54, "Raphael, The School of Athens, 1509–1511",
 "File__Raphael_The_School_of_A..._.190412", [
  ("Q1. Who is the artist?", "Raphael (Raffaello Sanzio da Urbino), Stanza della Segnatura, Vatican, 1509–1511."),
  ("Q2. What does it represent?", "An idealized gathering of the great philosophers of antiquity in a magnificent classical architectural setting. At centre: Plato (pointing upward, symbolizing the realm of Ideas) and Aristotle (pointing downward, symbolizing earthly knowledge)."),
  ("Q3. What High Renaissance characteristics are evident?", "The paradigmatic image of Renaissance humanism: celebration of classical learning, reason, and philosophy. Perfect one-point perspective; monumental classical architecture; harmonious, balanced composition; clarity and rational order. Raphael included portraits of contemporaries (Leonardo as Plato, Michelangelo as Heraclitus) among the ancient philosophers."),
  ("Q4. What Renaissance characteristics can be identified in this fresco?", "(1) Humanism — celebration of ancient Greek philosophers as sources of wisdom; (2) Ideal balance and harmony in composition; (3) Linear perspective creating deep architectural space; (4) Classical inspiration — the architecture resembles ancient Greek and Roman buildings; (5) Individualized psychological portraits of the philosophers."),
]),

(55, "Raphael, Sistine Madonna, 1512",
 "File_Raphael_Sistine_Madonna_.190410", [
  ("Q1. Who is the artist?", "Raphael, 1512. Gemäldegalerie Alte Meister, Dresden."),
  ("Q2. What does it represent?", "The Virgin Mary holds the Christ child; St. Sixtus (Pope) gestures toward the viewer while St. Barbara looks down. Two famous putti (cherubs) gaze upward at the bottom."),
  ("Q3. What characteristics are evident?", "Raphael perfects the Renaissance ideal of serene, harmonious beauty. The pyramidal composition is effortlessly natural. Raphael's achievement: making the divine seem naturally, perfectly human — the synthesis of Classical ideal beauty with Christian spiritual grace."),
]),

(56, "Raphael, Portrait of Pope Julius II, c. 1511–1512",
 "File__Raphael_Pope_Julius_II_.190411", [
  ("Q1. Who is the artist?", "Raphael, c. 1511–1512. National Gallery, London."),
  ("Q2. What is its significance?", "Considered the first psychologically penetrating papal portrait. Julius appears not as a triumphant symbol of power but as a weary, contemplative human being — his downward gaze and slightly hunched posture suggest the burdens of his office. This psychological realism follows Leonardo's innovations."),
]),

# ── VENETIAN SCHOOL ──────────────────────────────────────────────────────────
("VENETIAN SCHOOL (c. 1490–1600)", None, None, None),

(57, "Giorgione, Pastoral Symphony (Fête champêtre), c. 1508–1509",
 "File__Giorgione_Pastoral_Symp..._.191078", [
  ("Q1. Who is the artist?", "Attributed to Giorgione, with possible contributions by Titian, c. 1508–1509. Louvre, Paris."),
  ("Q2. What does it represent?", "Two clothed young men (possibly musicians) sit in a landscape; two nude women are with them — one holds a flute, one pours water. No narrative action — a lyrical, meditative 'mood painting.'"),
  ("Q3. What Venetian characteristics are evident?", "This work invented the 'pastoral' genre: lyrical, non-narrative scenes of figures in beautiful landscapes evoking poetry, music, and love. Key Venetian characteristics: primacy of colour (rich, warm tonalities) over Florentine line; landscape as active mood-setting element; soft, luminous light."),
]),

(58, "Giorgione, Sleeping Venus (Dresden Venus), c. 1508–1510",
 "File_Giorgione_Venus_asleep_.191079", [
  ("Q1. Who is the artist?", "Giorgione began this work c. 1508; may have been completed by Titian after Giorgione's death in 1510. Gemäldegalerie, Dresden."),
  ("Q2. What is its significance?", "Established the iconic 'reclining nude' type in Western painting that would be endlessly repeated (Titian's Venus of Urbino, Velázquez's Rokeby Venus, Manet's Olympia). The figure merges organically with the landscape — the curves of her body echo the hills behind her."),
]),

(59, "Titian, Venus of Urbino, 1538",
 "File_Titian_Venus_of_Urbino_.191080", [
  ("Q1. Who is the artist?", "Titian (Tiziano Vecellio), 1538. Uffizi Gallery, Florence."),
  ("Q2. How does it differ from Giorgione's Sleeping Venus?", "Giorgione: sleeping, outdoors, merging with nature, unaware of the viewer. Titian: awake, indoors, making direct eye contact — transforming the sleeping ideal into a sensual, self-possessed presence. This painting directly influenced Manet's provocative Olympia (1865)."),
  ("Q3. What Venetian characteristics are evident?", "Titian's virtuoso warm skin tones and luxurious textures (silk, fur) create sumptuous sensuality. Colour rather than line defines form — the defining characteristic of the Venetian school."),
]),

(60, "Titian, Bacchus and Ariadne, 1520–1523",
 "File__Titian_Bacchus_and_Ariadne_.191081", [
  ("Q1. Who is the artist?", "Titian, 1520–1523. National Gallery, London."),
  ("Q2. What does it represent?", "Bacchus leaps from his chariot, having spotted the abandoned Ariadne on Naxos. He falls in love with her and will transform her into a constellation (visible upper left). His riotous, intoxicated retinue follows."),
  ("Q3. What Venetian characteristics are evident?", "One of the great celebrations of Venetian colorism: brilliant ultramarine blue sky, vivid reds, glowing flesh tones — extraordinary visual richness. The dynamic composition — Bacchus mid-leap, Ariadne turning in surprise, the swirling retinue — captures energy and movement anticipating the Baroque."),
]),

(61, "Titian, The Flaying of Marsyas, c. 1570–1576",
 "File__Titian_Flaying_of_Marsyas_.191082", [
  ("Q1. Who is the artist?", "Titian, c. 1570–1576. Archbishop's Palace, Kroměříž, Czech Republic."),
  ("Q2. What does it represent?", "From Ovid's Metamorphoses: the satyr Marsyas challenged Apollo to a music contest and lost. Punishment: being flayed alive. Marsyas hangs upside down while Apollo skins him."),
  ("Q3. What characterizes Titian's late style?", "Loose, gestural brushwork that dissolves forms into flickering light and colour — surfaces built with multiple layers applied with fingers, brushes, and rags. Anticipates Impressionist and Expressionist techniques. The dark, meditative subject reflects old age and approaching death."),
]),

(62, "Titian, The Rape of Europa, 1560–1562",
 "File__Titian_The_rape_of_Europa_.191083", [
  ("Q1. Who is the artist?", "Titian, 1560–1562. Isabella Stewart Gardner Museum, Boston. One of his 'poesie' for Philip II of Spain."),
  ("Q2. What does it represent?", "Zeus (as white bull) abducts Europa, swimming out to sea while she holds his horn in terror and excitement. Cupids and sea creatures accompany the scene."),
  ("Q3. How does this compare to Boucher's treatment of the same myth?", "Titian (Venetian Renaissance): dynamic, turbulent; Europa's body tilted in dramatic agitation; sensual energy, the terror of divine power. Boucher (Rococo): pastel colours, playful atmosphere; Europa appears delighted rather than frightened; the violence and power of the myth suppressed entirely in favour of decorative pleasure."),
]),

(63, "Tintoretto, The Last Supper, 1592–1594",
 "File_Tintoretto_The_Last_Supper_.191085", [
  ("Q1. Who is the artist?", "Jacopo Tintoretto, 1592–1594. San Giorgio Maggiore, Venice."),
  ("Q2. How does it differ from Leonardo's Last Supper?", "Leonardo (High Renaissance, pre-Reformation): calm, symmetrical, horizontal, all figures clearly visible, rational clarity — humanist art confident in reason. Tintoretto (Baroque/Mannerist, post-Counter-Reformation): diagonal table rushing into deep space, chaotic atmosphere, ordinary servants mixed with disciples, supernatural angels materialising in smoke, dramatic chiaroscuro — emotional overwhelm rather than rational clarity."),
]),

(64, "Tintoretto, The Miracle of the Slave, 1548",
 "File_Tintoretto_The_miracle_o..._.191086", [
  ("Q1. Who is the artist?", "Tintoretto, 1548. Gallerie dell'Accademia, Venice."),
  ("Q2. What does it represent?", "St. Mark plunges headfirst from the sky to rescue a Christian slave about to be tortured by a pagan master. The instruments of torture break miraculously against the slave's body."),
  ("Q3. What characteristics are evident?", "Tintoretto combines Titian's colours with the dynamic figures of Michelangelo. The foreshortened St. Mark diving from above is dramatically dynamic. Many figures in varied poses and brilliant colours create enormous visual energy."),
]),

(65, "Tintoretto, Christ at the Sea of Galilee, c. 1575–1580",
 "File_Tintoretto_Christ_at_the..._.191087", [
  ("Q1. Who is the artist?", "Tintoretto, c. 1575–1580. National Gallery of Art, Washington D.C."),
  ("Q2. What does it represent?", "Christ walks on water toward the boat where the disciples have been struggling against a storm. The dramatic sea and stormy sky dwarf the human figures."),
  ("Q3. What Venetian/Mannerist characteristics are evident?", "The turbulent sea and storm-lit sky create overwhelming dramatic atmosphere. Unlike calm High Renaissance compositions, Tintoretto emphasises the power of nature and divine intervention over the small human figures. Dramatic chiaroscuro and swirling composition."),
]),

# ── MANNERISM ────────────────────────────────────────────────────────────────
("MANNERISM (c. 1520–1600)", None, None, None),

(66, "El Greco, The Burial of Count Orgaz, 1586–1588",
 "File__El_Greco_The_burial_of_..._.191089", [
  ("Q1. Who is the artist?", "El Greco (Doménikos Theotokópoulos), 1586–1588. Church of Santo Tomé, Toledo, Spain."),
  ("Q2. What does it represent?", "Lower half: the burial of Don Gonzalo Ruiz, Count of Orgaz, in which saints Stephen and Augustine miraculously appeared to lay his body in the grave. Upper half: Heaven, where Christ receives the count's soul, surrounded by the Virgin Mary, John the Baptist, and saints."),
  ("Q3. What Mannerist characteristics are evident?", "The painting is divided into two zones: realistic earthly scene (bottom) and heavenly vision above — a Mannerist distortion of natural space. El Greco's characteristic elongated figures, impossibly narrow bodies, and spiritual otherworldly atmosphere define Mannerism. Cool, supernatural light — not natural sunlight. Abandonment of Renaissance mathematical perspective in favour of spiritual hierarchy."),
]),

(67, "El Greco, Pietà (Laocoön), c. 1571–1576 / 1610–1614",
 "File_El_Greco_Laocoon_and_his..._.191088", [
  ("Q1. Who is the artist?", "El Greco (Doménikos Theotokópoulos). Mannerist period."),
  ("Q2. What Mannerist characteristics are evident?", "Elongated, twisting figures that defy natural proportions; cool, luminous, unnatural palette; compressed, shallow space pushing all figures to the foreground; tense, emotionally overwrought atmosphere. Compare with Michelangelo's Pietà (High Renaissance): serene, balanced, classical. El Greco's: anguished, distorted, spiritually intense — the Mannerist departure from Renaissance harmony."),
]),

(68, "Bronzino, Venus, Cupid, Folly and Time (Allegory), c. 1545",
 "File_Bronzino_Venus_cupid_and..._.193761", [
  ("Q1. Who is the artist?", "Agnolo Bronzino, c. 1545. National Gallery, London."),
  ("Q2. What does it represent?", "A complex erotic allegory: Venus and Cupid in an incestuous embrace, surrounded by Folly (throwing rose petals), Time (pulling back a blue curtain), Deceit (beautiful face but serpent's tail), and allegorical figures of jealousy, pleasure, and pain."),
  ("Q3. What Court Mannerist characteristics are evident?", "Quintessential Court Mannerism: extremely sophisticated, erotic, intellectually allusive work for a private learned patron. Characteristics: pale, cold, porcelain-like flesh; distorted, impossibly graceful poses (figura serpentinata); shallow compressed space; deliberately obscure allegorical meaning decodable only by the educated. Style over content, elegance over emotion."),
]),

(69, "Veronese, Banquet at the House of Levi, 1573",
 "File_Veronese_Banquet_at_the_..._.193762", [
  ("Q1. Who is the artist?", "Paolo Veronese (Paolo Caliari), 1573. Gallerie dell'Accademia, Venice. Originally titled The Last Supper but renamed after Inquisition objections."),
  ("Q2. What is its historical significance?", "Veronese was summoned before the Inquisition to explain the 'indecorous' elements (dwarfs, dogs, drunken soldiers). His defence — that painters have the same licence as poets to interpret subjects — anticipates modern ideas of artistic freedom. Rather than repaint, he simply changed the title."),
]),

# ── BAROQUE ──────────────────────────────────────────────────────────────────
("BAROQUE (c. 1600–1750)", None, None, None),

(70, "Bernini, David, 1623–1624 (marble)",
 "File_Bernini_David_.471310", [
  ("Q1. Who is the artist?", "Gian Lorenzo Bernini, 1623–1624. Borghese Gallery, Rome."),
  ("Q2. What does it represent?", "David in the act of slinging the stone at Goliath — the moment of maximum physical and psychological tension. The body is mid-coil, the face grimaces with fierce concentration."),
  ("Q3. How does it compare to Donatello's and Michelangelo's David?", "Donatello (Early Renaissance): calm, post-victory, contemplative. Michelangelo (High Renaissance): pre-battle tension, idealized composure. Bernini (Baroque): the violent action itself — body mid-twist, face grimacing, arm breaking into viewer's space, requiring the viewer to imagine Goliath on the other side. Baroque: dynamic movement, psychological intensity, figure engaging the viewer's space."),
]),

(71, "Caravaggio, The Calling of Saint Matthew, 1599–1600",
 "File__Caravaggio_The_calling_..._.194501", [
  ("Q1. Who is the artist?", "Michelangelo Merisi da Caravaggio, 1599–1600. Contarelli Chapel, San Luigi dei Francesi, Rome."),
  ("Q2. What does it represent?", "Matthew (a tax collector) sits at a table counting money. Christ and Peter enter from the right, Christ's arm outstretched. Matthew points to himself in disbelief: 'Me?' The divine calling is set in a contemporary Roman tavern."),
  ("Q3. What Baroque characteristics are evident?", "Caravaggio invented tenebrism: extreme chiaroscuro where figures emerge from near-total darkness, lit by a single sharp light source — creating theatrical drama, psychological intensity, and a sense of divine intrusion into the everyday world. The setting is a contemporary Roman tavern, not a biblical landscape — radical secularisation of sacred subjects."),
]),

(72, "Caravaggio, The Death of the Virgin, c. 1601–1606",
 "File_Caravaggio_The_death_of_..._.195557", [
  ("Q1. Who is the artist?", "Caravaggio, c. 1601–1606. Louvre, Paris. The painting was rejected by its patrons as too realistic."),
  ("Q2. What Baroque characteristics are evident?", "The Virgin appears as a common woman — her body bloated, feet bare. The apostles are rough, working-class men. No heavenly glory, no angels — just human loss. This radical naturalism brought sacred subjects into unidealized everyday human reality — Caravaggio's signature innovation. Tenebrism: deep darkness, dramatic light."),
]),

(73, "Caravaggio, Narcissus, c. 1594–1596",
 "File_Caravaggio_Narcissus_1594_.195540", [
  ("Q1. Who is the artist?", "Caravaggio, c. 1594–1596. Galleria Nazionale d'Arte Antica, Rome."),
  ("Q2. What does it represent?", "From Ovid: Narcissus kneels at a pool, gazing at his own reflection, in love with his image. His arms form a circle completed by the reflection below — formally embodying the mythological theme of self-absorption."),
  ("Q3. What characteristics are evident?", "Powerful tenebrism: the youth is lit from above against near-total darkness; the reflection in the dark water creates perfect symmetry. One of the earliest representations of the Narcissus myth in painting. Caravaggio creates profound psychological meaning through minimal composition and dramatic lighting."),
]),

(74, "Artemisia Gentileschi, Judith Slaying Holofernes, c. 1614–1620",
 "File_Jentileschi_Judith_slayi..._.194502", [
  ("Q1. Who is the artist?", "Artemisia Gentileschi, c. 1614–1620. Uffizi Gallery, Florence."),
  ("Q2. What does it represent?", "The biblical widow Judith beheads the Assyrian general Holofernes. Gentileschi depicts the act with unsparing violence — blood spurts as Judith and her maidservant pin Holofernes down and sever his head."),
  ("Q3. What Baroque characteristics are evident? Why is it significant?", "Caravaggio's tenebrism fully at work: dramatic light/shadow contrasts, figures emerging from darkness, violence lit with theatrical clarity. Gentileschi surpasses Caravaggio's own version in force and conviction — Judith performs the act with determined physical effort, not revulsion. Gentileschi is one of the first women to achieve recognition as a major artist."),
]),

(75, "Rubens, The Raising of the Cross, 1610–1611",
 "File_Rubens_The_rape_of_the_d..._.194503", [
  ("Q1. Who is the artist?", "Peter Paul Rubens, 1610–1611. Cathedral of Our Lady, Antwerp, Belgium."),
  ("Q2. What does it represent?", "Soldiers and workers strain to erect the cross bearing the crucified Christ, his white body a powerful diagonal against the dark sky and forest."),
  ("Q3. What Baroque characteristics are evident?", "(1) Dynamic diagonal composition; (2) Dramatic chiaroscuro — Christ's white body lit against dark figures and sky; (3) Physical energy and strain recalling the Laocoön; (4) Theatrical scale; (5) Drama of the moment — not High Renaissance serenity but physical effort and spiritual crisis."),
  ("Q4. What two Baroque techniques are evident?", "(1) Dynamic and dramatic composition — the diagonal arrangement of the cross and the straining muscular figures create intense physical and emotional energy; (2) Chiaroscuro — Christ's white body is dramatically illuminated against the dark, stormy sky, creating a theatrical stage-lighting effect."),
]),

(76, "Rubens, The Rape of the Daughters of Leucippus, c. 1617–1618",
 "File_Rubens_The_rape_of_the_d..._.194503", [
  ("Q1. Who is the artist?", "Rubens, c. 1617–1618. Alte Pinakothek, Munich."),
  ("Q2. What does it represent?", "The Dioscuri (Castor and Pollux) abduct the daughters of Leucippus. The two nude women are being lifted onto horses by the two armoured young men."),
  ("Q3. What Baroque characteristics are evident?", "Rubens's characteristic style: voluptuous, large-bodied female nudes with warm, glowing flesh; swirling, centrifugal composition of intertwined figures and horses; dynamic energy of bodies in motion."),
]),

(77, "Diego Velázquez, Las Meninas, 1656",
 "File__Diego_Velasquez_Los_Men..._.194504", [
  ("Q1. Who is the artist?", "Diego Velázquez, 1656. Museo del Prado, Madrid."),
  ("Q2. What does it represent?", "The Infanta Margarita surrounded by her maids of honour, a dwarf, a dog, and court figures. Velázquez appears at the left, working on a large canvas. In the background, a mirror reflects the blurry images of King Philip IV and Queen Mariana."),
  ("Q3. What Baroque characteristics and significance?", "Perhaps the most complex painting in Western art — it systematically confuses the roles of viewer, painter, model, and painted image. This meditation on painting, representation, and illusion is uniquely Baroque in its theatricality and self-consciousness. Foucault devoted the opening pages of 'The Order of Things' to it."),
]),

(78, "Diego Velázquez, Venus with a Mirror (Rokeby Venus), c. 1647–1651",
 "File__Diego_Velazquez_Venus_w..._.194505", [
  ("Q1. Who is the artist?", "Diego Velázquez, c. 1647–1651. National Gallery, London."),
  ("Q2. What is its significance?", "The only surviving female nude in Velázquez's work, one of very few in Spanish painting (heavily censored by the Inquisition). Engages in a subtle game: we see Venus from behind (private, vulnerable) yet her reflection looks back at us (aware, composed). The blurred, vague reflection — not showing what a mirror 'should' — is a deliberate painterly device."),
]),

(79, "Rembrandt van Rijn, The Blinding of Samson, 1636",
 "File_Rembrandt_The_Blinding_o..._.195227", [
  ("Q1. Who is the artist?", "Rembrandt van Rijn, 1636. Städel Museum, Frankfurt."),
  ("Q2. What does it represent?", "Delilah flees with Samson's cut hair while Philistine soldiers pin Samson down and one drives a dagger into his eye. A scene of extreme violence from the Book of Judges."),
  ("Q3. What Baroque characteristics are evident?", "Rembrandt is the Dutch master of tenebrism and psychological complexity. The scene is lit by blazing light from the tent opening, throwing the violent struggle into high contrast. Unlike Italian Baroque artists, Rembrandt uses light not only for drama but for profound spiritual meaning."),
]),

(80, "Rembrandt van Rijn, The Return of the Prodigal Son, c. 1668",
 "File_Rembrandt_The_return_of_..._.195665", [
  ("Q1. Who is the artist?", "Rembrandt van Rijn, c. 1668. Hermitage Museum, St. Petersburg."),
  ("Q2. What does it represent?", "From the Gospel of Luke: the prodigal son, who squandered his inheritance and returned penniless, kneels before his aged father who embraces him tenderly."),
  ("Q3. What characteristics are evident?", "One of the most profoundly moving paintings in Western art. Rembrandt's late style: figures emerge from golden-brown darkness; warm, intimate light concentrated on the father's hands and son's shaved head. The father's hands (one appears masculine, one more gentle) are among the most expressive details in painting. Rembrandt's genius: making divine love visible in ordinary human tenderness."),
]),

(81, "Poussin, The Abduction of the Sabine Women, c. 1633–1634",
 "File__Poussin_The_abduction_o..._.194506", [
  ("Q1. Who is the artist?", "Nicolas Poussin, c. 1633–1634. Metropolitan Museum of Art, New York."),
  ("Q2. What does it represent?", "Romulus signals his soldiers to seize the Sabine women at a festival — the founding event of Rome."),
  ("Q3. What French Classical Baroque characteristics are evident?", "Poussin represents French Baroque Classicism — not Rubens's swirling dynamism. Despite the chaotic scene, Poussin organises with clear, rational compositional structure: figures in balanced groups, architecture providing a stable grid. Drama is present but controlled. Figures have a sculptural, marmoreal quality."),
]),

(82, "Poussin, Et in Arcadia Ego, c. 1637–1638",
 "File_Poussin_Et_in_Arcadia_ego_.194507", [
  ("Q1. Who is the artist?", "Nicolas Poussin, c. 1637–1638. Louvre, Paris."),
  ("Q2. What does it represent?", "Three shepherds and a woman in the idyllic landscape of Arcadia read an inscription on a tomb: 'Et in Arcadia Ego' — 'Even in Arcadia, I (Death) am present.' The discovery of a tomb in paradise introduces the awareness of mortality."),
  ("Q3. What is its significance?", "One of the most philosophically meditated images in Western art. The shepherds try to decipher the inscription just as we do. Poussin's Classical style: clear, balanced composition; serene golden light; figures with the dignity of antique sculpture; a landscape of perfect ordered beauty. The tone is contemplative, elegiac — not dramatic."),
]),

# ── ROCOCO ───────────────────────────────────────────────────────────────────
("ROCOCO (c. 1720–1780)", None, None, None),

(83, "Fragonard, The Swing, 1767",
 "File__Fragonard_The_Swing_.195383", [
  ("Q1. Who is the artist?", "Jean-Honoré Fragonard, 1767. Wallace Collection, London."),
  ("Q2. What does it represent?", "A young woman swings in a lush garden while an older man (her husband) pushes her from behind; her young lover reclines in the bushes below, looking up her skirts. She kicks a shoe toward a statue of a cupid."),
  ("Q3. What Rococo characteristics are evident?", "The quintessential Rococo painting: (1) Playful, erotic subject matter — courtly love games, aristocratic leisure; (2) Light, feathery pastel palette — pink, pale blue, soft greens; (3) Soft, curving forms throughout; (4) Informal, asymmetrical composition; (5) Frivolous, unconcerned attitude — the world is pleasure and play; (6) Aristocratic, intimate private setting. Rococo was the art of French aristocracy immediately before the Revolution."),
]),

(84, "Fragonard, The Stolen Kiss, c. 1787–1788",
 "File__Fragonard_The_stolen_kiss_.195385", [
  ("Q1. Who is the artist?", "Fragonard, c. 1787–1788. Hermitage Museum, St. Petersburg."),
  ("Q2. What Rococo characteristics are evident?", "Light, intimate subject matter; soft warm palette; graceful flowing lines of the woman's dress; a sense of playful transgression and romantic excitement. The domestic intimacy reflects the Rococo interest in private life, small pleasures, and personal emotion."),
]),

(85, "Boucher, The Toilet of Venus (Bath of Venus), c. 1751",
 "File__Boucher_The_toilet_of_V..._.195386", [
  ("Q1. Who is the artist?", "François Boucher, c. 1751, commissioned by Madame de Pompadour. Metropolitan Museum of Art, New York."),
  ("Q2. What does it represent?", "Venus adorns herself after bathing, attended by playful cupids who bring her jewellery and doves. Set in a lush outdoor setting combining nature and luxury."),
  ("Q3. What Rococo characteristics are evident?", "Boucher is the master of Rococo sensuality and decorative elegance: pearly, porcelain-smooth skin of Venus; pastel palette of pink, blue, and gold; playful cupids; lush ornamental setting of silks, pearls, and flowers. Unlike Titian's earthy sensuality, Boucher's eroticism is light, decorative, and untroubled — pleasure without consequence."),
  ("Q4. How does this Rococo work contrast with David's Neoclassical Death of Marat?", "Boucher (Rococo): pastel palette, soft sensual forms, pleasurable subject, decorative and frivolous — art celebrating pleasure for pleasure's sake. David (Neoclassicism): stark dark palette, austere composition, moral and political subject (a revolutionary martyr), serious and didactic — art as moral and civic duty."),
]),

(86, "Boucher, The Rape/Abduction of Europa, 1734 & 1747",
 "File__Boucher_The_rape_of_Eur..._.195388", [
  ("Q1. Who is the artist?", "François Boucher, two versions: 1734 (Wallace Collection, London) and 1747 (Louvre, Paris)."),
  ("Q2. How does Boucher's treatment compare to Titian's and Rubens's?", "Titian (Venetian Renaissance): dynamic, turbulent sea, Europa tilted in terror and excitement, overwhelming divine power. Rubens (Baroque): large dynamic bodies, violent energy. Boucher (Rococo): lightness, pastel colours, decorative prettiness, Europa appears charmed and happy. The violent myth is drained of all dramatic content in favour of visual pleasure."),
]),

# ── NEOCLASSICISM ────────────────────────────────────────────────────────────
("NEOCLASSICISM (c. 1780–1820)", None, None, None),

(87, "Jacques-Louis David, The Death of Marat, 1793",
 "File_David_The_death_of_Marat_.195390", [
  ("Q1. Who is the artist?", "Jacques-Louis David, 1793. Musées Royaux des Beaux-Arts, Brussels."),
  ("Q2. What does it represent?", "Jean-Paul Marat, radical journalist of the French Revolution, has been assassinated in his medicinal bath by Charlotte Corday. Shown just after death — arm hanging with a quill, a letter in hand, wound visible in his chest. Inscribed on the crate: 'À Marat, David.'"),
  ("Q3. What Neoclassical characteristics are evident?", "David deliberately posed Marat to evoke Christian martyrdom — the drooping arm echoes Michelangelo's Dead Christ. Neoclassical characteristics: stark simplicity (the upper half of the painting is bare dark wall — no Rococo decoration); severe, undistracted composition; moral seriousness; the figure treated with sculptural dignity recalling classical reliefs."),
  ("Q4. What art movement does this belong to and what characteristics confirm it?", "Neoclassicism. Confirmed by: (1) austere, stripped-down composition with no decorative elements; (2) sculptural treatment of the figure — pale, marmoreal quality of ancient marble reliefs; (3) stoic emotional register — death shown with dignity and moral gravity, not melodrama; (4) clear, rational structure; (5) moral-political purpose — Neoclassicism believed art should serve virtue and civic duty."),
]),

(88, "Jacques-Louis David, The Death of Socrates, 1787",
 "File__David_Napoleon_crossing..._.195392", [
  ("Q1. Who is the artist?", "Jacques-Louis David, 1787. Metropolitan Museum of Art, New York."),
  ("Q2. What does it represent?", "From Plato's Phaedo: Socrates, condemned to death for impiety, drinks hemlock surrounded by grieving disciples. He gestures upward toward the realm of Ideas, accepting death with philosophical serenity."),
  ("Q3. What Neoclassical characteristics are evident?", "(1) Classical subject from ancient history/philosophy; (2) Moral didacticism — a model of heroic virtue, stoic acceptance of death for principles; (3) Sculptural, frieze-like arrangement of figures; (4) Clear, uncluttered composition with cool, even lighting (opposed to Baroque tenebrism); (5) Somber dignity rather than emotional excess."),
]),

(89, "Jacques-Louis David, Napoleon Crossing the Alps, 1800–1801",
 "File__David_Napoleon_crossing..._.195392", [
  ("Q1. Who is the artist?", "Jacques-Louis David, 1800–1801. Château de Malmaison, France."),
  ("Q2. What does it represent?", "Napoleon Bonaparte on a rearing horse crossing the Alps. He points forward toward his military objective. Inscribed on the rocks: 'Bonaparte,' 'Hannibal,' and 'Karolus Magnus' — linking Napoleon to great military leaders of history."),
  ("Q3. What Neoclassical characteristics are evident?", "Napoleon presented as the inheritor of Hannibal and Charlemagne — a historical hero of classical stature. (In reality, Napoleon crossed on a mule; he requested David paint him 'calm upon a fiery steed.') Neoclassical characteristics: heroic idealisation; reference to classical and historical precedent; formal, monumental composition; the figure as moral and political ideal."),
]),

# ── 19TH CENTURY ─────────────────────────────────────────────────────────────
("19TH CENTURY — ADDITIONAL", None, None, None),

(90, "Karl Briullov, The Last Day of Pompeii, 1830–1833",
 "File__Brullov_Karl_The_last_d..._.195394", [
  ("Q1. Who is the artist?", "Karl Briullov (Karl Pavlovich Bryullov), 1830–1833. Russian Museum, St. Petersburg."),
  ("Q2. What does it represent?", "The destruction of the Roman city of Pompeii by the eruption of Mount Vesuvius in 79 CE. Citizens flee as ash, lightning, and falling masonry destroys the city."),
  ("Q3. What stylistic characteristics are evident?", "Blends academic Neoclassicism (idealised sculptural figures, clear compositional structure) with Romantic elements (dramatic catastrophic subject matter, spectacular lighting effects, emotional intensity, the sublime power of nature destroying civilisation). A sensation in Europe when exhibited — praised by Pushkin and Walter Scott."),
]),

]  # end ARTWORKS


# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = section.bottom_margin = Inches(1)
    section.left_margin   = section.right_margin  = Inches(1)

# title
t = doc.add_heading("Art History — Complete Study Guide", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Ancient Greece · Renaissance · Mannerism · Baroque · Rococo · Neoclassicism")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

for entry in ARTWORKS:
    if entry[1] is None:
        # section header
        doc.add_heading(entry[0], level=1)
        continue

    num, title, folder, qas = entry

    add_artwork_title(doc, num, title)

    # image
    img_path = find_image(folder) if folder else None
    if img_path:
        add_image(doc, img_path)
    else:
        p = doc.add_paragraph(f"[Image not found — check folder: {folder}]")
        p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        p.runs[0].font.italic = True

    # Q&A
    for q, a in qas:
        add_qa(doc, q, a)

    doc.add_paragraph("")  # spacer


out = "Art_History_Study_Guide.docx"
doc.save(out)
print(f"\n✅  Done!  Saved as: {out}")
print(f"   Artworks processed: {sum(1 for e in ARTWORKS if e[1] is not None)}")